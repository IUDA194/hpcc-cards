from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from main import table

# Загрузка документа
table_ret = table()  # Создание экземпляра класса Table
returned_dict = table_ret.get_data()  # Вызов метода get_data() на экземпляре table
print(returned_dict)



doc = Document("ref.docx") #editing font/fontsize
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(14)
# Ваш словарь с данными (замените на реальные данные)
data_dict = {#filling first page
    # "Галузь знань –": "Галузь знань – " + returned_dict['occupation'],
    # "Спеціальність -": "Спеціальність - " + returned_dict['spaciallity'],
    # "Освітньо-професійна програма -": "Освітньо-професійна програма - " + returned_dict['programm'],
    # "Освітньо-професійний ступень -": "Освітньо-професійний ступень - " + returned_dict['ladder'],
    # "Форма здобуття освіти -": "Форма здобуття освіти - " + returned_dict['form']
    "Галузь знань –": f"Галузь знань – {returned_dict.get('occupation')}",#не работает ибо не понятен ключ по которому искать это.
    "Спеціальність -": f"Спеціальність - {returned_dict.get('spaciallity')}",
    "Освітньо-професійна програма -": f"Освітньо-професійна програма - {returned_dict.get('programm')}",
    "Освітньо-професійний ступень -": f"Освітньо-професійний ступень - {returned_dict.get('ladder')}",
    "Форма здобуття освіти -": f"Форма здобуття освіти - {returned_dict.get('form')}"
}

# Проход по всем параграфам в документе
for paragraph in doc.paragraphs:
    for field, value in data_dict.items():
        if field in paragraph.text:
            paragraph.text = paragraph.text.replace(field, value)

data_dict = {#заполнение некст странницы
   # 1: {"Вид практики": "значение1", "Кількість кредитів ЄКТС": "значение2", "Курс": "значение3"},
    2: {"Вид практики": "значение4", "Кількість кредитів ЄКТС": "значение5", "Курс": "значение6", "Від": "29231", "До": "24112", "Дата захисту": "228", "Оцінка за шкалою закладу освіти": "5", "Прізвище викладача": "Абоба", "123": "fas"},
    3: {"Вид практики": "значение4", "Кількість кредитів ЄКТС": "значение5", "Курс": "значение6", "Від": "29231", "До": "24112", "Дата захисту": "228", "Оцінка за шкалою закладу освіти": "5", "Прізвище викладача": "Абоба", "123": "fas"},
    4: {"Вид практики": "значение4", "Кількість кредитів ЄКТС": "значение5", "Курс": "значение6", "Від": "29231", "До": "24112", "Дата захисту": "228", "Оцінка за шкалою закладу освіти": "5", "Прізвище викладача": "Абоба", "123": "fas"},
}

# Номер строки в таблице, с которой начинаются данные
start_row = 1

# Проход по строкам таблицы и вставка данных
for row_num, row_data in data_dict.items():
    table = doc.tables[0]  # Первая таблица в документе
    row = table.rows[start_row + row_num - 1].cells
    
    # Вставка данных из словаря в соответствующие ячейки таблицы
    for col_num, value in enumerate(row_data.values()):
        row[col_num].text = str(value)

data_dict_table2 = {
    # 1: {"Освітній компонент": "значение1", "Кількість кредитів ЄКТС": "значение2", "Вид (КП/КР)": "значение3", ...},
    2: {"Освітній компонент": "значение4", "Кількість кредитів ЄКТС": "значение5", "Вид (КП/КР)": "значение6",},
    # Добавьте остальные данные в формате аналогичном выше
}

# Номер строки во второй таблице, с которой начинаются данные
start_row_table2 = 1

# Проход по строкам второй таблицы и вставка данных
for row_num, row_data in data_dict_table2.items():
    table = doc.tables[1]  # Вторая таблица в документе
    row = table.rows[start_row_table2 + row_num - 1].cells
    
    # Вставка данных из словаря в соответствующие ячейки таблицы
    for col_num, value in enumerate(row_data.values()):
        row[col_num].text = str(value)
        
data_dict_text = {
    "Тема кваліфікаційної роботи –": "новое_значение1",
    "Прізвище керівника –": "новое_значение2",
    "Дата здачі закінченої кваліфікаційної роботи –": "новое_значение3",
    "Дата захисту роботи –": "новое_значение4",
    "Оцінка екзаменаційної комісії –": "новое_значение5"
}

# Проход по всем параграфам в документе
for paragraph in doc.paragraphs:
    for field, value in data_dict_text.items():
        if field in paragraph.text:
            paragraph.text = paragraph.text.replace(field, value)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Выравнивание по левому краю
            

# Сохранение измененного документа
doc.save("save.docx")
