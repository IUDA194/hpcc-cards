from docx import Document

# Open the .docx file
doc = Document('path/to/your/file.docx')

# Access the content of the document
paragraphs = doc.paragraphs
tables = doc.tables

# Modify the content
for paragraph in paragraphs:
    if 'example' in paragraph.text:
        paragraph.text = paragraph.text.replace('example', 'replacement')

# Save the modified document
doc.save('path/to/save/modified/file.docx')
